
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from pdf2image import convert_from_path
import re
import os
import redis
import hashlib
from PIL import Image
import torch
import torchvision.transforms as T
from torchvision.transforms.functional import InterpolationMode
from transformers import AutoModel, AutoTokenizer

# === Redis client (for OCR only) ===
redis_client = redis.Redis.from_url("redis://localhost:6379")

def get_ocr_cache(image_path: str) -> str | None:
    with open(image_path, "rb") as f:
        file_bytes = f.read()
        content_hash = hashlib.md5(file_bytes).hexdigest()
    key = "ocr:" + content_hash + ":" + os.path.basename(image_path)
    result = redis_client.get(key)
    if result:
        print("✅ [OCR CACHE HIT]", image_path)
        return result.decode("utf-8")
    return None

def set_ocr_cache(image_path: str, value: str):
    with open(image_path, "rb") as f:
        file_bytes = f.read()
        content_hash = hashlib.md5(file_bytes).hexdigest()
    key = "ocr:" + content_hash + ":" + os.path.basename(image_path)
    redis_client.setex(key, 21600, value)
    print("💾 [OCR CACHE SAVE]", image_path)

# ==== MÔ HÌNH CHAT ==== #
from langchain_community.cache import RedisCache as BaseRedisCache
from langchain.globals import set_llm_cache

class RedisCache(BaseRedisCache):
    def lookup(self, prompt: str, llm_string: str):
        result = super().lookup(prompt, llm_string)
        if result is not None:
            print("✅ [CACHE HIT] Câu trả lời đã được lấy từ Redis cache.")
        return result

    def update(self, prompt: str, llm_string: str, return_val: str):
        print("💾 [CACHE SAVE] Lưu kết quả mới vào Redis cache.")
        return super().update(prompt, llm_string, return_val)

set_llm_cache(RedisCache(redis_client, ttl=21600))

# ==== VINTERN OCR ==== #
_model_vintern = None
_tokenizer_vintern = None
_transform = T.Compose([
    T.Lambda(lambda img: img.convert("RGB") if img.mode != "RGB" else img),
    T.Resize((448, 448), interpolation=InterpolationMode.BICUBIC),
    T.ToTensor(),
    T.Normalize(mean=(0.485, 0.456, 0.406), std=(0.229, 0.224, 0.225)),
])

def init_vintern():
    global _model_vintern, _tokenizer_vintern
    if _model_vintern is None or _tokenizer_vintern is None:
        print("🔄 Đang tải mô hình Vintern-1B-v3_5 từ HuggingFace...")
        _model_vintern = AutoModel.from_pretrained(
            "5CD-AI/Vintern-1B-v3_5",
            torch_dtype=torch.bfloat16 if torch.cuda.is_available() else torch.float32,
            low_cpu_mem_usage=True,
            trust_remote_code=True
        ).eval()
        _model_vintern.to(torch.device("cuda" if torch.cuda.is_available() else "cpu"))

        _tokenizer_vintern = AutoTokenizer.from_pretrained(
            "5CD-AI/Vintern-1B-v3_5", trust_remote_code=True, use_fast=False
        )
        print("Đã tải Vintern-1B-v3_5.")

def vintern_ocr_cached(image_path: str) -> str:
    init_vintern()

    cached_result = get_ocr_cache(image_path)
    if cached_result:
        return cached_result

    print("🔍 [OCR RUNNING]", image_path)
    image = Image.open(image_path).convert("RGB")
    pixel_values = _transform(image).unsqueeze(0)

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    dtype = torch.bfloat16 if torch.cuda.is_available() else torch.float32
    pixel_values = pixel_values.to(device=device, dtype=dtype)
    _model_vintern.to(device)

    prompt = "<image>\nMô tả hình ảnh một cách chi tiết trả về dạng markdown."
    raw_response = _model_vintern.chat(
        _tokenizer_vintern,
        pixel_values,
        prompt,
        generation_config={"max_new_tokens": 2048, "do_sample": False, "num_beams": 3, "repetition_penalty": 3.5}
    )

    if isinstance(raw_response, list) and isinstance(raw_response[0], dict) and "generated_text" in raw_response[0]:
        response = raw_response[0]["generated_text"]
    elif isinstance(raw_response, str):
        response = raw_response
    else:
        response = str(raw_response)

    lines = response.strip().splitlines()
    seen = set()
    filtered = []
    for line in lines:
        if line not in seen:
            filtered.append(line)
            seen.add(line)
    response = "\n".join(filtered)
    set_ocr_cache(image_path, response)
    return response

# def get_answer(prompt: str, messages: list, vector_store) -> str:
#     messages.append(HumanMessage(content=prompt))
#     docs = vector_store.similarity_search(prompt, k=5)
#     context = "\n\n".join([doc.page_content for doc in docs])
#     rag_prompt = f"""
#     Bạn là một trợ lý AI chính xác, chỉ dựa trên nội dung sau để trả lời câu hỏi.

#     ### Nội dung tài liệu:
#     {context}

#     ### Câu hỏi:
#     {prompt}
#     ### Hướng dẫn:
#         - Nếu tìm thấy thông tin, hãy trả lời ngắn gọn, đúng trọng tâm.
#         - Nếu không tìm thấy thông tin trong tài liệu, hãy trả lời: "Không tìm thấy thông tin liên quan trong tài liệu."
#     """
#     llm = ChatOllama(model="vistral-7b-chat", temperature=0.7, base_url="http://localhost:11434")
#     full_response = ""
#     for chunk in llm.stream(messages[:-1] + [HumanMessage(content=rag_prompt)]):
#         if chunk.content:
#             print(chunk.content, end="", flush=True)
#             full_response += chunk.content
#     print()
#     messages.append(AIMessage(content=full_response))
#     return full_response

def get_answer(question: str, messages: list, vector_store) -> str:
    docs = vector_store.similarity_search(question, k=5)
    context = "\n\n".join([doc.page_content for doc in docs])

    if not context.strip():
        return "Không tìm thấy thông tin liên quan trong tài liệu."

    prompt = f"""
        Bạn là một trợ lý AI chính xác, chỉ dựa trên nội dung sau để trả lời câu hỏi.

        ### Nội dung tài liệu:
        {context}

        ### Câu hỏi:
        {question}

        ### Hướng dẫn:
        - Trả lời ngắn gọn, đúng trọng tâm.
        - Nếu không tìm thấy thông tin, hãy nói rõ điều đó.
    """

    messages.append(HumanMessage(content=prompt))
    llm = ChatOllama(model="vistral-7b-chat", temperature=0.7, base_url="http://localhost:11434")

    full_response = ""
    for chunk in llm.stream(messages):
        if chunk.content:
            print(chunk.content, end="", flush=True)
            full_response += chunk.content
    print()
    messages.append(AIMessage(content=full_response))
    return full_response


def initialize_embeddings():
    return HuggingFaceEmbeddings(
        model_name="BAAI/bge-m3",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

def load_vector_store():
    return FAISS.load_local("vectorstores/db_faiss", embeddings=initialize_embeddings(), allow_dangerous_deserialization=True)

def create_temp_vectorstore(text: str):
    splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
    docs = splitter.split_documents([Document(page_content=text)])
    return FAISS.from_documents(docs, initialize_embeddings())

def read_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def remove_repeated_phrases(text: str, threshold: int = 5) -> str:
    phrases = re.findall(r'(.{5,100}?)(?=(\s\1)+)', text)
    cleaned = text
    for phrase, _ in phrases:
        if cleaned.count(phrase) > threshold:
            cleaned = cleaned.replace(phrase, '', cleaned.count(phrase) - threshold)
    return cleaned

def process_file_content(file_path: str) -> str:
    fname = file_path.lower()
    if fname.endswith(".docx"):
        raw = read_docx(file_path)
    elif fname.endswith(".pdf"):
        images = convert_from_path(file_path, poppler_path=r"poppler-24.08.0\Library\bin")
        all_contents = []
        for idx, img in enumerate(images):
            tmp_path = f"temp_{idx}.png"
            img.save(tmp_path)
            content = vintern_ocr_cached(tmp_path)
            all_contents.append(f"[Trang {idx+1}]\n{content}")
            os.remove(tmp_path)
        raw = "\n\n".join(all_contents)
    elif fname.endswith((".png", ".jpg", ".jpeg")):
        raw = vintern_ocr_cached(file_path)
    else:
        raise ValueError("Định dạng file không được hỗ trợ")

    return remove_repeated_phrases(raw)

def initialize_chat_messages():
    return [SystemMessage(content="Bạn là một trợ lý AI thông minh và hữu ích. Hãy hỗ trợ người dùng với các câu hỏi và tài liệu.")]

def print_welcome_message():
    print("=== CHATBOT KHỞI ĐỘNG ===")
    print("Gõ 'exit' để thoát, 'clear' để xóa lịch sử chat")
    print("Gõ 'file:đường_dẫn_tệp' để hỏi từ file (.pdf, .docx, .png, .jpg)")
    print("===========================")

def run_chat():
    vector_store = load_vector_store()
    messages = initialize_chat_messages()
    print_welcome_message()

    while True:
        user_input = input("Bạn: ").strip()

        if user_input.lower() == "exit":
            print("=== KẾT THÚC CHAT ===")
            break

        if user_input.lower() == "clear":
            messages = initialize_chat_messages()
            print("Đã xóa lịch sử chat!")
            continue

        if user_input.startswith("file:"):
            file_path = user_input[5:].strip()
            if not os.path.exists(file_path):
                print("Không tìm thấy tệp!")
                continue
            try:
                content = process_file_content(file_path)
                print(f"📄 Nội dung trích xuất từ {file_path}:")
                print(content[:1500], "...\n")
                while True:
                    q = input("Câu hỏi về tệp (enter để bỏ qua): ").strip()
                    if not q:
                        break
                    temp_vs = create_temp_vectorstore(content)
                    context_docs = temp_vs.similarity_search(q, k=3)
                    context = "\n\n".join([doc.page_content for doc in context_docs])
                    full_prompt = f"Bối cảnh từ nội dung tệp:\n{context}\n\nCâu hỏi: {q}"
                    print("AI: ", end="")
                    get_answer(full_prompt, messages, vector_store)
            except Exception as e:
                print(f"Lỗi xử lý file: {e}")
            continue

        print("AI: ", end="")
        get_answer(user_input, messages, vector_store)

def main():
    run_chat()

if __name__ == "__main__":
    main()
